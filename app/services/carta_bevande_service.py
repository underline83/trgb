# @version: v1.0-fase3-export
# -*- coding: utf-8 -*-
"""
TRGB — Service Carta Bevande

Costruisce HTML/PDF/DOCX per la Carta delle Bevande completa:
aperitivi, birre, vini (delega a carta_vini_service), amari fatti in casa,
amari & liquori, distillati, tisane, tè.

Funzioni pure: non fa I/O disco, non gestisce richieste HTTP.
Le letture DB sono isolate in `_load_voci_attive(sezione_key)` e
`_load_sezioni_attive()`: restano qui per tenere il router thin.

Riferimento design: docs/carta_bevande_design.md
"""

from __future__ import annotations

import json
from datetime import datetime
from itertools import groupby
from typing import Any, Iterable, Optional

from app.models.bevande_db import (
    count_voci_by_sezione,
    get_bevande_conn,
    get_version_timestamp,
    list_sezioni,
)


# ────────────────────────────────────────────────────────────
# VERSIONE
# ────────────────────────────────────────────────────────────

def get_version_string() -> str:
    """
    Ritorna 'v{YYYY}.{MM}.{DD}' basata su MAX(updated_at) delle tabelle
    bevande. Se il DB è vuoto o senza timestamp, usa la data corrente.

    Non è un semver "vero", è una versione umana leggibile nel footer
    del PDF che riflette l'ultima modifica significativa alla carta.
    """
    ts = get_version_timestamp()
    if not ts:
        now = datetime.now()
        return f"v{now.year}.{now.month:02d}.{now.day:02d}"
    try:
        # updated_at è 'YYYY-MM-DD HH:MM:SS' (datetime('now','localtime'))
        dt = datetime.strptime(ts[:19], "%Y-%m-%d %H:%M:%S")
    except ValueError:
        try:
            dt = datetime.strptime(ts[:10], "%Y-%m-%d")
        except ValueError:
            dt = datetime.now()
    return f"v{dt.year}.{dt.month:02d}.{dt.day:02d}"


# ────────────────────────────────────────────────────────────
# HELPER — LOAD DATA
# ────────────────────────────────────────────────────────────

def _row_to_dict(row) -> dict[str, Any]:
    if row is None:
        return {}
    d = dict(row)
    for field in ("schema_form", "tags", "extra"):
        if field in d and d[field]:
            try:
                d[field] = json.loads(d[field])
            except (TypeError, ValueError):
                pass
    return d


def _load_sezioni_attive() -> list[dict[str, Any]]:
    return [_row_to_dict(s) for s in list_sezioni(only_active=True)]


def _load_voci_attive(sezione_key: str) -> list[dict[str, Any]]:
    conn = get_bevande_conn()
    try:
        rows = conn.execute(
            """
            SELECT *
              FROM bevande_voci
             WHERE sezione_key = ? AND attivo = 1
             ORDER BY ordine, id
            """,
            (sezione_key,),
        ).fetchall()
        return [_row_to_dict(r) for r in rows]
    finally:
        conn.close()


# ────────────────────────────────────────────────────────────
# HELPER — FORMATTERS
# ────────────────────────────────────────────────────────────

def _format_prezzo(voce: dict[str, Any]) -> str:
    """Regole: prezzo_label ha precedenza. Altrimenti prezzo_eur → '€ XX,YY'. Vuoto se nulla."""
    label = voce.get("prezzo_label")
    if label:
        return str(label)
    eur = voce.get("prezzo_eur")
    if eur in (None, "", 0, 0.0):
        return ""
    try:
        return f"€ {float(eur):.2f}".replace(".", ",")
    except (TypeError, ValueError):
        return str(eur)


def _esc(s: Any) -> str:
    """Escape minimo per evitare HTML injection accidentale nei campi utente."""
    if s is None:
        return ""
    return (
        str(s)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def _note_staff_block(voce: dict[str, Any], staff: bool) -> str:
    if not staff:
        return ""
    note = voce.get("note_interne")
    if not note:
        return ""
    return f"<div class='bev-note-staff'>🔒 {_esc(note)}</div>"


# ────────────────────────────────────────────────────────────
# RENDER — PATTERN A: tabella_4col (distillati, amari_liquori)
# ────────────────────────────────────────────────────────────

def _render_tabella_4col(voci: list[dict[str, Any]], staff: bool) -> str:
    """
    Tabella compatta: Tipologia (raggruppato) · Produttore · Nome · Prezzo.
    Per distillati uso 'tipologia' (Grappa, Rum, Whisky…); per amari_liquori
    non c'è una tipologia richiesta, quindi se è NULL raggruppo sotto 'Amari & Liquori'.
    """
    if not voci:
        return "<p class='bev-empty'><em>Nessuna voce attiva.</em></p>"

    def k_tip(v): return v.get("tipologia") or "—"

    # Groupby richiede lista ordinata; voci già ordinata per ordine,id → non è
    # la chiave di tipologia, quindi ordino esplicitamente per stabilità.
    voci_sorted = sorted(voci, key=lambda v: (k_tip(v), v.get("ordine") or 0, v.get("id") or 0))

    out = ["<div class='bev-4col'>"]
    for tip, g in groupby(voci_sorted, k_tip):
        g = list(g)
        show_tip_header = tip and tip != "—"
        if show_tip_header:
            out.append(f"<h3 class='bev-4col-group'>{_esc(tip)}</h3>")

        out.append("<table class='bev-4col-table'><tbody>")
        for v in g:
            produttore = v.get("produttore") or ""
            nome = v.get("nome") or ""
            regione = v.get("regione") or ""
            # Colonna di contesto: regione per distillati (molto usata), vuota se manca
            ctx = regione
            prezzo = _format_prezzo(v)
            out.append(
                "<tr>"
                f"<td class='b4-ctx'>{_esc(ctx)}</td>"
                f"<td class='b4-prod'>{_esc(produttore)}</td>"
                f"<td class='b4-nome'>{_esc(nome)}</td>"
                f"<td class='b4-prezzo'>{_esc(prezzo)}</td>"
                "</tr>"
            )
            # Descrizione inline (se presente) in riga separata piccola
            desc = v.get("descrizione")
            if desc:
                out.append(
                    "<tr class='b4-desc-row'>"
                    f"<td colspan='4' class='b4-desc'>{_esc(desc)}</td>"
                    "</tr>"
                )
            note_html = _note_staff_block(v, staff)
            if note_html:
                out.append(
                    f"<tr class='b4-staff-row'><td colspan='4'>{note_html}</td></tr>"
                )
        out.append("</tbody></table>")
    out.append("</div>")
    return "".join(out)


# ────────────────────────────────────────────────────────────
# RENDER — PATTERN B: scheda_estesa (birre, aperitivi, amari_casa)
# ────────────────────────────────────────────────────────────

def _render_scheda_estesa(voci: list[dict[str, Any]], staff: bool) -> str:
    """
    Scheda estesa: nome grande, riga meta (stile/formato/grad/IBU), descrizione,
    prezzo allineato a destra. Pensato per birre con più attributi.
    """
    if not voci:
        return "<p class='bev-empty'><em>Nessuna voce attiva.</em></p>"

    out = ["<div class='bev-scheda'>"]
    for v in voci:
        nome = v.get("nome") or ""
        sottotit = v.get("sottotitolo") or ""
        produttore = v.get("produttore") or ""
        formato = v.get("formato") or ""
        grad = v.get("gradazione")
        ibu = v.get("ibu")
        desc = v.get("descrizione") or ""
        prezzo = _format_prezzo(v)

        # Riga meta: pezzi disponibili, separati da ·
        meta_parts = []
        if produttore:
            meta_parts.append(_esc(produttore))
        if sottotit:
            meta_parts.append(f"<em>{_esc(sottotit)}</em>")
        if formato:
            meta_parts.append(_esc(formato))
        if grad not in (None, "", 0, 0.0):
            try:
                meta_parts.append(f"{float(grad):.1f}%".replace(".", ","))
            except (TypeError, ValueError):
                meta_parts.append(_esc(grad))
        if ibu not in (None, "", 0):
            try:
                meta_parts.append(f"IBU {int(ibu)}")
            except (TypeError, ValueError):
                meta_parts.append(f"IBU {_esc(ibu)}")
        meta_line = " · ".join(meta_parts)

        out.append("<div class='bev-scheda-item'>")
        out.append(
            "<div class='bev-scheda-head'>"
            f"<span class='bev-scheda-nome'>{_esc(nome)}</span>"
            f"<span class='bev-scheda-prezzo'>{_esc(prezzo)}</span>"
            "</div>"
        )
        if meta_line:
            out.append(f"<div class='bev-scheda-meta'>{meta_line}</div>")
        if desc:
            out.append(f"<div class='bev-scheda-desc'>{_esc(desc)}</div>")
        staff_html = _note_staff_block(v, staff)
        if staff_html:
            out.append(staff_html)
        out.append("</div>")

    out.append("</div>")
    return "".join(out)


# ────────────────────────────────────────────────────────────
# RENDER — PATTERN C: nome_badge_desc (tisane, tè)
# ────────────────────────────────────────────────────────────

_BADGE_COLORS_TE = {
    "nero":    ("#5b3a29", "#ffffff"),
    "verde":   ("#5a7a3a", "#ffffff"),
    "oolong":  ("#b36b2a", "#ffffff"),
    "rosso":   ("#a33c2a", "#ffffff"),
    "puer":    ("#3a2e24", "#ffffff"),
    "bianco":  ("#e9e3d4", "#2b2118"),
    "tisana":  ("#7a6b5a", "#ffffff"),
}


def _badge_html(tipologia: str | None, sottotit: str | None) -> str:
    """
    Per tè: badge colorato sul tipo.
    Per tisane: badge grigio sulla categoria (sottotitolo).
    """
    if tipologia:
        key = (tipologia or "").strip().lower()
        bg, fg = _BADGE_COLORS_TE.get(key, ("#7a6b5a", "#ffffff"))
        label = tipologia
        return (
            f"<span class='bev-badge' style=\"background:{bg};color:{fg}\">"
            f"{_esc(label)}</span>"
        )
    if sottotit:
        return f"<span class='bev-badge bev-badge-neutral'>{_esc(sottotit)}</span>"
    return ""


def _render_nome_badge_desc(voci: list[dict[str, Any]], staff: bool) -> str:
    """
    Pattern C: ogni voce è NOME + badge (tipologia colorata / categoria) +
    descrizione. Prezzo opzionale.
    """
    if not voci:
        return "<p class='bev-empty'><em>Nessuna voce attiva.</em></p>"

    out = ["<div class='bev-badgelist'>"]
    for v in voci:
        nome = v.get("nome") or ""
        tipologia = v.get("tipologia") or ""
        sottotit = v.get("sottotitolo") or ""
        paese = v.get("paese_origine") or ""
        desc = v.get("descrizione") or ""
        prezzo = _format_prezzo(v)

        badge = _badge_html(tipologia, sottotit)
        meta_parts = []
        if paese:
            meta_parts.append(_esc(paese))
        meta_line = " · ".join(meta_parts)

        out.append("<div class='bev-badge-item'>")
        out.append(
            "<div class='bev-badge-head'>"
            f"<span class='bev-badge-nome'>{_esc(nome)}</span>"
            f"{badge}"
            f"<span class='bev-badge-prezzo'>{_esc(prezzo)}</span>"
            "</div>"
        )
        if meta_line:
            out.append(f"<div class='bev-badge-meta'>{meta_line}</div>")
        if desc:
            out.append(f"<div class='bev-badge-desc'>{_esc(desc)}</div>")
        staff_html = _note_staff_block(v, staff)
        if staff_html:
            out.append(staff_html)
        out.append("</div>")

    out.append("</div>")
    return "".join(out)


# ────────────────────────────────────────────────────────────
# DISPATCHER — build_section_html
# ────────────────────────────────────────────────────────────

_LAYOUT_DISPATCH = {
    "tabella_4col":    _render_tabella_4col,
    "scheda_estesa":   _render_scheda_estesa,
    "nome_badge_desc": _render_nome_badge_desc,
}


def build_section_html(
    sezione: dict[str, Any],
    voci: list[dict[str, Any]],
    for_pdf: bool = False,
    staff: bool = False,
) -> str:
    """
    Rende una singola sezione della carta bevande.
    - Layout 'vini_dinamico': gestito dall'orchestratore (chiama carta_vini_service).
    - Altri layout: dispatch su _LAYOUT_DISPATCH.
    - 'for_pdf': se True, avvolge in .bev-section-pdf (page-break-before su sezione ≠ prima).
    """
    layout = sezione.get("layout") or "scheda_estesa"
    if layout == "vini_dinamico":
        # Non renderizzato qui: l'orchestratore inietta la carta vini separatamente.
        return ""

    render_fn = _LAYOUT_DISPATCH.get(layout, _render_scheda_estesa)
    body_html = render_fn(voci, staff=staff)

    intro_html = sezione.get("intro_html") or ""
    intro_block = f"<div class='bev-intro'>{intro_html}</div>" if intro_html else ""
    wrap_class = "bev-section bev-section-pdf" if for_pdf else "bev-section"
    # id='sez-<key>' serve come target per `target-counter` nel TOC (numeri pagina).
    sez_id = f"sez-{_esc(sezione.get('key') or '')}"
    return (
        f"<section id='{sez_id}' class='{wrap_class}' data-sezione='{_esc(sezione.get('key'))}'>"
        f"<h2 class='bev-section-title'>{_esc(sezione.get('nome'))}</h2>"
        f"{intro_block}"
        f"{body_html}"
        f"</section>"
    )


# ────────────────────────────────────────────────────────────
# COPERTINA + TOC MASTER
# ────────────────────────────────────────────────────────────

def build_copertina_html(logo_path: Optional[str] = None, staff: bool = False) -> str:
    """Frontespizio master per il PDF."""
    data_oggi = datetime.now().strftime("%d/%m/%Y")
    version = get_version_string()
    title = "CARTA DELLE BEVANDE — STAFF" if staff else "CARTA DELLE BEVANDE"
    logo_block = (
        f"<img src='file://{logo_path}' class='front-logo'>" if logo_path else ""
    )
    staff_tag = "<div class='front-subtitle'>VERSIONE INTERNA</div>" if staff else ""
    return f"""
    <div class="front-page">
        {logo_block}
        <div class="front-title">{title}</div>
        {staff_tag}
        <div class="front-date">Aggiornata al {data_oggi}</div>
        <div class="front-version">{version}</div>
    </div>
    """


def build_toc_html(sezioni: list[dict[str, Any]]) -> str:
    """
    Indice UNICO master della carta bevande.

    Regole:
    - Include solo sezioni con almeno 1 voce attiva (skip categorie vuote).
    - Rispetta l'ordine `ordine` delle sezioni (settabile dalle frecce admin
      nella sidebar della shell carta).
    - Per la sezione 'vini' espande inline il sotto-indice
      tipologia → nazione → regione, usando lo stesso builder della carta vini
      (`build_carta_toc_html`) ma senza il wrapper `<div class='toc-page'>`
      per non aprire un secondo indice annidato.
    - Stile (dal 2026-04-20, variante D.3): le macro-sezioni
      (Vini, Aperitivi, Amari di Casa, …) usano la classe `toc-macro`
      (18pt, peso 400, uppercase, tracking 0.32em, colore #5a4634),
      distinta dalle sotto-voci della carta vini che restano in
      `toc-tipologia` (14pt bold). Così la gerarchia macro/sub è
      leggibile a colpo d'occhio nell'indice PDF.
    """
    if not sezioni:
        return ""

    # Conteggi voci (per sezioni bevande standard)
    counts = count_voci_by_sezione()

    # Carica vini rows una sola volta (serve per: (a) capire se vini è "vuota",
    # (b) espandere il sub-indice).
    vini_rows: list[dict[str, Any]] = []
    try:
        from app.repositories.vini_repository import load_vini_ordinati
        vini_rows = list(load_vini_ordinati())
    except Exception:
        vini_rows = []

    rows_html: list[str] = ["<div class='toc-page'>", "<div class='toc-title'>INDICE</div>"]
    any_section = False

    for s in sezioni:
        nome = s.get("nome") or ""
        key = s.get("key") or ""

        if key == "vini":
            # Skip se cantina vuota
            if not vini_rows:
                continue
            any_section = True
            # Ancora verso <section id='sez-vini'> → target-counter per numero pagina.
            rows_html.append(
                f"<a class='toc-macro' href='#sez-vini'>"
                f"<span class='toc-name'>{_esc(nome)}</span>"
                f"<span class='toc-leader'></span>"
                f"<span class='toc-pn'></span>"
                f"</a>"
            )
            # Sub-indice tipologie/nazioni/regioni — stesso stile della carta vini.
            # Chiamiamo build_carta_toc_html e strippiamo il wrapper esterno
            # <div class='toc-page'> + <div class='toc-title'> per fondere
            # il sotto-indice nel master.
            try:
                from app.services.carta_vini_service import build_carta_toc_html
                vini_toc_raw = build_carta_toc_html(vini_rows)
                # Rimuovi <div class='toc-page'>…INDICE…</div> wrapper:
                # build_carta_toc_html emette "<div class='toc-page'><div class='toc-title'>INDICE</div>…</div>"
                # Tagliamo il wrapper esterno.
                prefix = "<div class='toc-page'><div class='toc-title'>INDICE</div>"
                suffix = "</div>"
                inner = vini_toc_raw
                if inner.startswith(prefix) and inner.endswith(suffix):
                    inner = inner[len(prefix):-len(suffix)]
                rows_html.append(inner)
            except Exception:
                # Fallback silenzioso: in peggio restiamo col solo titolo sezione
                pass
            rows_html.append("<div class='toc-spacer'></div>")
            continue

        # Sezioni bevande standard (aperitivi, birre, amari_casa, ...)
        c = counts.get(key, {"totale": 0, "attive": 0})
        if int(c.get("attive", 0) or 0) <= 0:
            continue  # salta sezioni vuote
        any_section = True
        # Ancora verso <section id='sez-<key>'> → target-counter per numero pagina.
        rows_html.append(
            f"<a class='toc-macro' href='#sez-{_esc(key)}'>"
            f"<span class='toc-name'>{_esc(nome)}</span>"
            f"<span class='toc-leader'></span>"
            f"<span class='toc-pn'></span>"
            f"</a>"
        )
        rows_html.append("<div class='toc-spacer'></div>")

    if not any_section:
        return ""
    rows_html.append("</div>")
    return "".join(rows_html)


# ────────────────────────────────────────────────────────────
# ORCHESTRATORE MASTER
# ────────────────────────────────────────────────────────────

def build_carta_bevande_html(
    include_vini: bool = True,
    for_pdf: bool = False,
    staff: bool = False,
) -> str:
    """
    Assembla il BODY HTML della carta completa.
    Non include <html>/<head>/<body> né CSS link: quelle restano a cura del router.

    Ordine: sezioni attive (sorted by ordine), rispettando:
    - 'vini_dinamico' → delega a carta_vini_service (nel suo stesso stile).
    - altre sezioni → build_section_html.
    """
    # Import locale per evitare circolarità in caso di reload dev.
    # build_carta_toc_html non serve piu' qui: il master TOC e' costruito da
    # build_toc_html (questo service) che delega a carta_vini_service per il
    # sub-indice tipologie/nazioni/regioni.
    from app.services.carta_vini_service import (
        build_carta_body_html,
        build_carta_body_html_htmlsafe,
        build_calici_section_html,
        build_calici_section_htmlsafe,
    )
    from app.repositories.vini_repository import (
        load_vini_ordinati,
        load_vini_calici,
    )

    sezioni = _load_sezioni_attive()
    parts: list[str] = []

    for s in sezioni:
        key = s.get("key")
        layout = s.get("layout") or "scheda_estesa"

        if key == "vini" and layout == "vini_dinamico":
            if not include_vini:
                continue
            # Blocco vini completo (calici + corpo). NB: l'indice dettagliato
            # tipologia/nazione/zona NON va piu' qui: e' gia' incluso nel
            # master TOC unico (build_toc_html lo espande inline).
            vini_rows = list(load_vini_ordinati())
            calici_rows = list(load_vini_calici())
            # Skip sezione vini se cantina vuota: evita <h2>Vini</h2> seguito da
            # "Nessun vino da mostrare" nel corpo carta (sessione 2026-04-20).
            if not vini_rows and not calici_rows:
                continue
            if for_pdf:
                calici_html = build_calici_section_html(calici_rows)
                body_vini = build_carta_body_html(vini_rows)
                parts.append(
                    "<section id='sez-vini' class='bev-section bev-section-pdf bev-section-vini'>"
                    "<h2 class='bev-section-title'>Vini</h2>"
                    f"<div class='carta-body'>{calici_html}{body_vini}</div>"
                    "</section>"
                )
            else:
                calici_html = build_calici_section_htmlsafe(calici_rows)
                body_vini = build_carta_body_html_htmlsafe(vini_rows)
                parts.append(
                    "<section id='sez-vini' class='bev-section bev-section-vini'>"
                    "<h2 class='bev-section-title'>Vini</h2>"
                    f"{calici_html}"
                    f"{body_vini}"
                    "</section>"
                )
            continue

        # Sezioni bevande standard
        voci = _load_voci_attive(key) if key else []
        # Skip sezione senza voci attive: evita blocchi "Tisane" (o altro) con
        # solo titolo e nulla sotto (sessione 2026-04-20).
        if not voci:
            continue
        parts.append(build_section_html(s, voci, for_pdf=for_pdf, staff=staff))

    return "".join(parts)


# ────────────────────────────────────────────────────────────
# DOCX
# ────────────────────────────────────────────────────────────

def build_carta_bevande_docx(logo_path=None, staff: bool = False) -> "Document":
    """
    Genera un oggetto Document python-docx con la carta bevande completa.
    Usa paragrafi + tabelle senza bordi (pattern analogo a build_carta_docx vini).
    """
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    from app.services.carta_vini_service import build_carta_docx
    from app.repositories.vini_repository import load_vini_ordinati

    doc = Document()

    # Margini
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # ── Helper tabella senza bordi ──
    def _no_borders(cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        borders = tc.makeelement(qn("w:tcBorders"), {})
        for edge in ("top", "left", "bottom", "right"):
            el = tc.makeelement(qn(f"w:{edge}"), {
                qn("w:val"): "none",
                qn("w:sz"): "0",
                qn("w:space"): "0",
                qn("w:color"): "auto",
            })
            borders.append(el)
        tcPr.append(borders)

    # ── Copertina ──
    if logo_path and hasattr(logo_path, "exists") and logo_path.exists():
        doc.add_picture(str(logo_path), width=Inches(1.8))

    data_oggi = datetime.now().strftime("%d/%m/%Y")
    version = get_version_string()
    title = "CARTA DELLE BEVANDE — STAFF" if staff else "CARTA DELLE BEVANDE"
    doc.add_heading(title, level=0)
    p_sub = doc.add_paragraph(f"Osteria Tre Gobbi — Aggiornata al {data_oggi}  ·  {version}")
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sezioni = _load_sezioni_attive()

    for s in sezioni:
        key = s.get("key") or ""
        layout = s.get("layout") or "scheda_estesa"
        nome = s.get("nome") or key

        doc.add_page_break()
        doc.add_heading(nome, level=1)

        if s.get("intro_html"):
            # Molto grezzo: strip tag HTML dall'intro per il DOCX
            import re as _re
            intro_txt = _re.sub(r"<[^>]+>", "", s["intro_html"]).strip()
            if intro_txt:
                p = doc.add_paragraph(intro_txt)
                p.paragraph_format.space_after = Pt(8)

        if key == "vini" and layout == "vini_dinamico":
            # Delega al builder vini ma montato nel documento corrente sarebbe
            # complicato: più semplice creare un sub-doc e importarne il corpo.
            # Per MVP: rimandiamo a "vedi carta vini separata" nel DOCX; il DOCX
            # completo vini si scarica da /vini/carta/docx.
            p = doc.add_paragraph(
                "La sezione Vini è disponibile come documento dedicato. "
                "Scarica la Carta Vini completa da Menu → Vini → Carta → Word."
            )
            p.paragraph_format.space_after = Pt(8)
            continue

        voci = _load_voci_attive(key)
        if not voci:
            doc.add_paragraph("(Nessuna voce attiva)")
            continue

        if layout == "tabella_4col":
            # Tabella compatta: Tipologia · Produttore · Nome · Prezzo
            tbl = doc.add_table(rows=0, cols=4)
            tbl.autofit = True
            for v in voci:
                row = tbl.add_row().cells
                for c in row:
                    _no_borders(c)
                row[0].text = v.get("tipologia") or v.get("regione") or ""
                row[1].text = v.get("produttore") or ""
                row[2].text = v.get("nome") or ""
                row[3].text = _format_prezzo(v)
                row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for c in row:
                    for r_ in c.paragraphs[0].runs:
                        r_.font.size = Pt(10)
                if v.get("descrizione"):
                    p = doc.add_paragraph(v["descrizione"])
                    p.paragraph_format.space_after = Pt(4)
                    for r_ in p.runs:
                        r_.font.size = Pt(9)
                        r_.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                if staff and v.get("note_interne"):
                    p = doc.add_paragraph(f"[STAFF] {v['note_interne']}")
                    for r_ in p.runs:
                        r_.font.size = Pt(9)
                        r_.italic = True
                        r_.font.color.rgb = RGBColor(0xA0, 0x40, 0x00)

        elif layout == "nome_badge_desc":
            for v in voci:
                p = doc.add_paragraph()
                rn = p.add_run(v.get("nome") or "")
                rn.bold = True
                rn.font.size = Pt(12)
                badge_txt = v.get("tipologia") or v.get("sottotitolo") or ""
                if badge_txt:
                    rb = p.add_run(f"  [{badge_txt}]")
                    rb.italic = True
                    rb.font.size = Pt(9)
                    rb.font.color.rgb = RGBColor(0x5B, 0x2C, 0x1A)
                prezzo = _format_prezzo(v)
                if prezzo:
                    rp = p.add_run(f"   {prezzo}")
                    rp.bold = True
                    rp.font.size = Pt(10)
                if v.get("descrizione"):
                    pd = doc.add_paragraph(v["descrizione"])
                    pd.paragraph_format.space_after = Pt(4)
                    for r_ in pd.runs:
                        r_.font.size = Pt(9)
                if staff and v.get("note_interne"):
                    ps = doc.add_paragraph(f"[STAFF] {v['note_interne']}")
                    for r_ in ps.runs:
                        r_.font.size = Pt(9)
                        r_.italic = True
                        r_.font.color.rgb = RGBColor(0xA0, 0x40, 0x00)

        else:  # scheda_estesa (default)
            for v in voci:
                p = doc.add_paragraph()
                rn = p.add_run(v.get("nome") or "")
                rn.bold = True
                rn.font.size = Pt(12)
                prezzo = _format_prezzo(v)
                if prezzo:
                    p.add_run("   ")
                    rp = p.add_run(prezzo)
                    rp.bold = True
                    rp.font.size = Pt(10)

                meta_bits = []
                if v.get("produttore"):
                    meta_bits.append(v["produttore"])
                if v.get("sottotitolo"):
                    meta_bits.append(v["sottotitolo"])
                if v.get("formato"):
                    meta_bits.append(v["formato"])
                if v.get("gradazione") not in (None, "", 0, 0.0):
                    try:
                        meta_bits.append(f"{float(v['gradazione']):.1f}%".replace(".", ","))
                    except (TypeError, ValueError):
                        meta_bits.append(str(v["gradazione"]))
                if v.get("ibu") not in (None, "", 0):
                    meta_bits.append(f"IBU {v['ibu']}")
                if meta_bits:
                    pm = doc.add_paragraph(" · ".join(str(x) for x in meta_bits))
                    pm.paragraph_format.space_after = Pt(2)
                    for r_ in pm.runs:
                        r_.italic = True
                        r_.font.size = Pt(9)
                        r_.font.color.rgb = RGBColor(0x5A, 0x46, 0x34)

                if v.get("descrizione"):
                    pd = doc.add_paragraph(v["descrizione"])
                    pd.paragraph_format.space_after = Pt(6)
                    for r_ in pd.runs:
                        r_.font.size = Pt(10)

                if staff and v.get("note_interne"):
                    ps = doc.add_paragraph(f"[STAFF] {v['note_interne']}")
                    for r_ in ps.runs:
                        r_.font.size = Pt(9)
                        r_.italic = True
                        r_.font.color.rgb = RGBColor(0xA0, 0x40, 0x00)

    return doc
