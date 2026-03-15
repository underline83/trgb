# @version: v1.1-carta-vini-service
# -*- coding: utf-8 -*-
"""
TRGB — Service Carta Vini

Responsabilità:
- Funzioni pure di costruzione HTML per:
  - corpo carta (PDF)
  - corpo carta (preview HTML safe)
  - indice (toc)

NON si occupa di:
- Lettura/scrittura file
- Endpoints FastAPI
- Gestione richieste
- Scelta paths su disco

Queste funzioni sono usate dal router vini.
"""

from __future__ import annotations
from itertools import groupby
from typing import Iterable, Dict, Any

import unicodedata
import re


# ------------------------------------------------------------
# UTILS CONDIVISE
# ------------------------------------------------------------

def slugify(value: str) -> str:
    """Crea un id CSS/HTML semplice per tipologie e regioni."""
    if not value:
        return "x"
    value_norm = unicodedata.normalize("NFKD", value)
    value_ascii = value_norm.encode("ascii", "ignore").decode("ascii")
    value_ascii = re.sub(r"[^a-zA-Z0-9]+", "-", value_ascii).strip("-").lower()
    return value_ascii or "x"


def resolve_regione(r: Dict[str, Any]) -> str:
    """Regola base: Regione → Nazione → Varie."""
    if r.get("REGIONE"):
        return r["REGIONE"]
    if r.get("NAZIONE"):
        return r["NAZIONE"]
    return "Varie"


# ------------------------------------------------------------
# BUILDER — CORPO CARTA (PDF)
# ------------------------------------------------------------
def build_carta_body_html(rows: Iterable[Dict[str, Any]]) -> str:
    rows = list(rows)
    if not rows:
        return "<p><em>Nessun vino da mostrare.</em></p>"

    def k_tip(r): return r["TIPOLOGIA"] or "Senza tipologia"
    def k_naz(r): return r.get("NAZIONE") or "Varie"
    def k_reg(r): return resolve_regione(r)
    def k_prod(r): return r["PRODUTTORE"] or "Produttore sconosciuto"

    html = ""

    for tip, g1 in groupby(rows, k_tip):
        g1 = list(g1)
        html += f"<h2 class='tipologia'>{tip}</h2>"

        for naz, g1b in groupby(g1, k_naz):
            g1b = list(g1b)
            html += (
                f"<div class='nazione'>"
                f"<span class='naz-line'></span>"
                f"<span class='naz-label'>{naz}</span>"
                f"<span class='naz-line'></span>"
                f"</div>"
            )

            for reg, g2 in groupby(g1b, k_reg):
                g2 = list(g2)
                html += f"<h4 class='regione'>{reg}</h4>"

                for prod, g3 in groupby(g2, k_prod):
                    g3 = list(g3)
                    html += "<div class='producer-block'>"
                    html += "<div class='spacer'></div>"
                    html += f"<h5 class='produttore'>{prod}</h5>"
                    html += "<table class='vini'><tbody>"

                    for r in g3:
                        desc = r["DESCRIZIONE"] or ""
                        annata = r["ANNATA"] or ""
                        prezzo = r["PREZZO"]
                        if prezzo not in (None, ""):
                            try:
                                prezzo = f"€ {float(prezzo):.2f}".replace(".", ",")
                            except Exception:
                                prezzo = str(prezzo)
                        else:
                            prezzo = ""

                        html += (
                            "<tr>"
                            f"<td class='vino'>{desc}</td>"
                            f"<td class='annata'>{annata}</td>"
                            f"<td class='prezzo'>{prezzo}</td>"
                            "</tr>"
                        )

                    html += "</tbody></table></div>"

    return html


# ------------------------------------------------------------
# BUILDER — HTML SAFE (PREVIEW)
# ------------------------------------------------------------
def build_carta_body_html_htmlsafe(rows: Iterable[Dict[str, Any]]) -> str:
    """Preview HTML identica alla 1.33 (senza producer-block)."""
    rows = list(rows)
    if not rows:
        return "<p><em>Nessun vino.</em></p>"

    def k_tip(r): return r["TIPOLOGIA"] or "Senza tipologia"
    def k_naz(r): return r.get("NAZIONE") or "Varie"
    def k_reg(r): return resolve_regione(r)
    def k_prod(r): return r["PRODUTTORE"] or "Produttore sconosciuto"

    html = ""

    for tip, g1 in groupby(rows, k_tip):
        g1 = list(g1)
        html += f"<h2 class='tipologia'>{tip}</h2>"

        for naz, g1b in groupby(g1, k_naz):
            g1b = list(g1b)
            html += (
                f"<div class='nazione'>"
                f"<span class='naz-line'></span>"
                f"<span class='naz-label'>{naz}</span>"
                f"<span class='naz-line'></span>"
                f"</div>"
            )

            for reg, g2 in groupby(g1b, k_reg):
                g2 = list(g2)
                html += f"<h4 class='regione'>{reg}</h4>"

                for prod, g3 in groupby(g2, k_prod):
                    g3 = list(g3)
                    html += "<div class='spacer'></div>"
                    html += f"<h5 class='produttore'>{prod}</h5>"
                    html += "<table class='vini'><tbody>"

                    for r in g3:
                        desc = r["DESCRIZIONE"] or ""
                        annata = r["ANNATA"] or ""
                        prezzo = r["PREZZO"]
                        if prezzo not in (None, ""):
                            try:
                                prezzo = f"€ {float(prezzo):.2f}".replace(".", ",")
                            except Exception:
                                prezzo = str(prezzo)

                        html += (
                            "<tr>"
                            f"<td class='vino'>{desc}</td>"
                            f"<td class='annata'>{annata}</td>"
                            f"<td class='prezzo'>{prezzo}</td>"
                            "</tr>"
                        )

                    html += "</tbody></table>"

    return html


# ------------------------------------------------------------
# BUILDER — INDICE
# ------------------------------------------------------------
def build_carta_toc_html(rows: Iterable[Dict[str, Any]]) -> str:
    rows = list(rows)
    if not rows:
        return ""

    def k_tip(r): return r["TIPOLOGIA"] or "Senza tipologia"
    def k_naz(r): return r.get("NAZIONE") or "Varie"
    def k_reg(r): return resolve_regione(r)

    html = [
        "<div class='toc-page'>",
        "<div class='toc-title'>INDICE</div>"
    ]

    for tip, g1 in groupby(rows, k_tip):
        g1 = list(g1)
        html.append(f"<div class='toc-tipologia'>{tip}</div>")

        for naz, g1b in groupby(g1, k_naz):
            g1b = list(g1b)
            html.append(f"<div class='toc-nazione'>· {naz}</div>")

            seen = set()
            for reg, g2 in groupby(g1b, k_reg):
                g2 = list(g2)
                if reg not in seen:
                    seen.add(reg)
                    html.append(f"<div class='toc-regione'>&nbsp;&nbsp;— {reg}</div>")

        html.append("<div class='toc-spacer'></div>")

    html.append("</div>")
    return "".join(html)


# ------------------------------------------------------------
# BUILDER — DOCX (condiviso tra /carta/docx e /cantina/docx)
# ------------------------------------------------------------
def build_carta_docx(rows: Iterable[Dict[str, Any]], logo_path=None) -> "Document":
    """
    Genera un oggetto Document python-docx con la carta vini completa.
    Usa tabelle senza bordi per allineare descrizione / annata / prezzo,
    come nell'HTML/PDF.
    """
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    rows = list(rows)
    doc = Document()

    # -- Stile base: margini pagina
    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # A4 content width = 21cm - 2cm - 2cm = 17cm
    CONTENT_WIDTH = Cm(17)
    COL_DESC = Cm(11.5)   # 67.6%
    COL_ANNATA = Cm(2.5)  # 14.7%
    COL_PREZZO = Cm(3.0)  # 17.6%

    def _no_borders(cell):
        """Rimuove bordi da una cella."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        borders = tcPr.find(qn("w:tcBorders"))
        if borders is not None:
            tcPr.remove(borders)
        borders = tc.makeelement(qn("w:tcBorders"), {})
        for edge in ("top", "left", "bottom", "right"):
            el = tc.makeelement(qn(f"w:{edge}"), {
                qn("w:val"): "none",
                qn("w:sz"): "0",
                qn("w:space"): "0",
                qn("w:color"): "auto",
            })
            borders.append(el)
        tcPr.append(borders)

    def _set_cell_width(cell, width):
        """Forza la larghezza di una cella."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = tcPr.find(qn("w:tcW"))
        if tcW is None:
            tcW = tc.makeelement(qn("w:tcW"), {})
            tcPr.append(tcW)
        tcW.set(qn("w:w"), str(int(width.emu / 635)))  # EMU → twips
        tcW.set(qn("w:type"), "dxa")

    def _cell_margins(cell, top=0, bottom=0, left=0, right=0):
        """Setta i margini interni della cella (in twips)."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        mar = tc.makeelement(qn("w:tcMar"), {})
        for name, val in [("top", top), ("bottom", bottom),
                          ("start", left), ("end", right)]:
            el = tc.makeelement(qn(f"w:{name}"), {
                qn("w:w"): str(val),
                qn("w:type"): "dxa",
            })
            mar.append(el)
        tcPr.append(mar)

    def _add_wine_table(doc, wines):
        """Aggiunge una tabella senza bordi con righe vino."""
        from docx.table import Table as DocxTable
        tbl = doc.add_table(rows=len(wines), cols=3)
        tbl.autofit = False

        # Larghezza colonne via XML
        tblW = tbl._tbl.find(qn("w:tblPr"))
        if tblW is None:
            tblW = tbl._tbl.makeelement(qn("w:tblPr"), {})
            tbl._tbl.insert(0, tblW)

        for i, wine in enumerate(wines):
            desc, annata, prezzo = wine
            row = tbl.rows[i]

            # Cella descrizione
            c0 = row.cells[0]
            _no_borders(c0)
            _set_cell_width(c0, COL_DESC)
            _cell_margins(c0, top=10, bottom=10, left=28, right=40)
            p0 = c0.paragraphs[0]
            p0.paragraph_format.space_before = Pt(0)
            p0.paragraph_format.space_after = Pt(0)
            r0 = p0.add_run(desc)
            r0.font.size = Pt(9)

            # Cella annata
            c1 = row.cells[1]
            _no_borders(c1)
            _set_cell_width(c1, COL_ANNATA)
            _cell_margins(c1, top=10, bottom=10, left=0, right=40)
            p1 = c1.paragraphs[0]
            p1.paragraph_format.space_before = Pt(0)
            p1.paragraph_format.space_after = Pt(0)
            p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r1 = p1.add_run(annata)
            r1.font.size = Pt(9)
            r1.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            # Cella prezzo
            c2 = row.cells[2]
            _no_borders(c2)
            _set_cell_width(c2, COL_PREZZO)
            _cell_margins(c2, top=10, bottom=10, left=0, right=28)
            p2 = c2.paragraphs[0]
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after = Pt(0)
            p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r2 = p2.add_run(prezzo)
            r2.font.size = Pt(9)
            r2.bold = True

        return tbl

    # -- Logo
    if logo_path and logo_path.exists():
        doc.add_picture(str(logo_path), width=Inches(1.8))

    # -- Titolo
    from datetime import datetime
    data_oggi = datetime.now().strftime("%d/%m/%Y")
    doc.add_heading("CARTA DEI VINI", level=0)
    doc.add_paragraph(f"Osteria Tre Gobbi — Aggiornata al {data_oggi}")

    if not rows:
        doc.add_paragraph("Nessun vino da mostrare.")
        return doc

    # -- Grouping keys
    def k_tip(r): return r["TIPOLOGIA"] or "Senza tipologia"
    def k_naz(r): return r.get("NAZIONE") or "Varie"
    def k_reg(r): return resolve_regione(r)
    def k_prod(r): return r["PRODUTTORE"] or "Produttore sconosciuto"

    for tip, g1 in groupby(rows, k_tip):
        g1 = list(g1)
        doc.add_heading(tip, level=1)

        for naz, g1b in groupby(g1, k_naz):
            g1b = list(g1b)
            doc.add_heading(naz, level=2)

            for reg, g2 in groupby(g1b, k_reg):
                g2 = list(g2)
                # Regione: corsivo + grassetto
                p_reg = doc.add_paragraph()
                run_reg = p_reg.add_run(reg)
                run_reg.italic = True
                run_reg.bold = True
                run_reg.font.size = Pt(11)

                for prod, g3 in groupby(g2, k_prod):
                    g3 = list(g3)
                    # Produttore: grassetto
                    p_prod = doc.add_paragraph()
                    p_prod.paragraph_format.space_before = Pt(6)
                    p_prod.paragraph_format.space_after = Pt(2)
                    rr = p_prod.add_run(prod)
                    rr.bold = True
                    rr.font.size = Pt(10)

                    # Prepara righe vino
                    wines = []
                    for r in g3:
                        desc = r["DESCRIZIONE"] or ""
                        annata = r["ANNATA"] or ""
                        prezzo = r["PREZZO"]
                        if prezzo not in (None, "", 0):
                            try:
                                prezzo = f"€ {float(prezzo):.2f}".replace(".", ",")
                            except Exception:
                                prezzo = str(prezzo)
                        else:
                            prezzo = ""
                        wines.append((desc, annata, prezzo))

                    # Tabella senza bordi
                    _add_wine_table(doc, wines)

    return doc
