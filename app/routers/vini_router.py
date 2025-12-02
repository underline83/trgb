# @version: v2.06-stable
# -*- coding: utf-8 -*-
"""
Router Vini — HTML + PDF + DOCX + Movimenti + Import Excel

Linea 2.x = motore carta avanzato (impaginazione + indice)

Changelog v2.06:
 - CLEAN: rimosso uso di tabella 'vini_raw'
 - CLEAN: /vini/upload ora scrive solo su 'vini' (vini.sqlite3)
 - REFINE: import pulito, niente import interni ridondanti

Changelog v2.05:
 - ADD: endpoint /vini/{vino_id}/movimenti (GET/POST)
        con log utente (JWT) nelle note movimento

Changelog v2.04:
 - ADD: endpoint /vini/carta/pdf-staff (PDF staff)
        (per ora stesso contenuto del PDF standard, con label STAFF)
 - KEEP: /vini/carta/pdf invariato per il cliente
 - KEEP: HTML preview identica alla linea 1.3x
"""

from __future__ import annotations

from itertools import groupby
from pathlib import Path
from datetime import datetime
import unicodedata
import re
from typing import List, Dict, Any

import os
import tempfile

import pandas as pd
from fastapi import APIRouter, UploadFile, File, HTTPException, Query, Depends
from fastapi.responses import HTMLResponse, JSONResponse, FileResponse
from pydantic import BaseModel, Field
from weasyprint import HTML, CSS
from docx import Document
from docx.shared import Inches

from app.services.carta_vini_service import (
    build_carta_body_html,
    build_carta_body_html_htmlsafe,
    build_carta_toc_html,
)
from app.services.auth_service import get_current_user

from app.models.vini_db import (
    get_connection,
    init_database,
    get_vino_by_id,
    list_movimenti_vino,
    registra_movimento,
)
from app.models.vini_model import (
    clear_vini_table,
    normalize_dataframe,
    insert_vini_rows,
)
from app.repositories.vini_repository import load_vini_ordinati


router = APIRouter(prefix="/vini", tags=["Vini"])

print(">>>>> LOADING VINI_ROUTER v2.06   PATH:", __file__)

# PATH DI BASE
BASE_DIR = Path(__file__).resolve().parents[2]
STATIC_DIR = BASE_DIR / "static"
CSS_HTML = STATIC_DIR / "css" / "carta_html.css"
CSS_PDF = STATIC_DIR / "css" / "carta_pdf.css"
LOGO_PATH = STATIC_DIR / "img" / "logo_tregobbi.png"


# ------------------------------------------------------------
# UTILS
# ------------------------------------------------------------
def slugify(value: str) -> str:
    """Crea un id CSS/HTML semplice per tipologie e regioni."""
    if not value:
        return "x"
    value_norm = unicodedata.normalize("NFKD", value)
    value_ascii = value_norm.encode("ascii", "ignore").decode("ascii")
    value_ascii = re.sub(r"[^a-zA-Z0-9]+", "-", value_ascii).strip("-").lower()
    return value_ascii or "x"


def resolve_regione(r):
    """Regola base: Regione → Nazione → Varie."""
    if r["REGIONE"]:
        return r["REGIONE"]
    if r["NAZIONE"]:
        return r["NAZIONE"]
    return "Varie"


# ------------------------------------------------------------
# UPLOAD EXCEL → DB 'vini'
# ------------------------------------------------------------
@router.post("/upload")
async def upload_vini(
    file: UploadFile = File(...),
    format: str = Query("json", enum=["json", "html"]),
):
    """
    Importa il foglio 'VINI' da un Excel, normalizza con vini_model
    e popola la tabella 'vini' nel DB ufficiale (vini.sqlite3).

    Nessun DB RAW: Excel → normalize_dataframe → insert_vini_rows.
    """
    # 1) Salvo il file temporaneamente
    with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp:
        tmp.write(await file.read())
        tmp_path = tmp.name

    try:
        # 2) Leggo il foglio "VINI"
        try:
            xls = pd.ExcelFile(tmp_path)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Errore apertura Excel: {e}")

        if "VINI" not in xls.sheet_names:
            raise HTTPException(status_code=400, detail="Foglio 'VINI' non trovato nel file Excel.")

        raw = pd.read_excel(xls, sheet_name="VINI")

        # 3) Pulizia righe completamente vuote (ma NON tocchiamo i nomi colonne)
        raw = raw.fillna("")
        raw = raw[raw.apply(lambda r: any(str(x).strip() for x in r), axis=1)]
        tot = len(raw)

        # 4) Normalizzazione verso lo schema 'vini'
        df = normalize_dataframe(raw)

        # 5) Scrittura nel DB ufficiale
        conn = get_connection()
        init_database()
        clear_vini_table(conn)
        inserite, errori, count = insert_vini_rows(conn, df)
        conn.close()

    finally:
        # 6) Provo a cancellare il file temporaneo
        try:
            os.remove(tmp_path)
        except Exception:
            pass

    # 7) Risposta HTML (come prima) o JSON
    if format == "html":
        max_val = max(count.values()) if count else 1
        html = "<html><body><h2>Risultato Import</h2>"
        html += f"<p>Totali Excel: {tot} — Inserite in DB: {inserite}</p>"
        html += "<table border='1' cellpadding='6'>"
        for t, c in sorted(count.items(), key=lambda x: (-x[1], x[0])):
            width = int((c / max_val) * 100)
            html += (
                f"<tr><td>{t}</td><td>{c}</td>"
                f"<td><div style='background:#90c490;height:14px;width:{width}%'></div></td></tr>"
            )
        html += "</table></body></html>"
        return HTMLResponse(html)

    return {
        "righe_totali_excel": tot,
        "righe_inserite_db": inserite,
        "errori": errori[:100],
    }


# ------------------------------------------------------------
# HTML PREVIEW CARTA
# ------------------------------------------------------------
@router.get("/carta", response_class=HTMLResponse)
def genera_carta_vini_html():
    rows = list(load_vini_ordinati())
    body = build_carta_body_html_htmlsafe(rows)

    html = f"""
    <html>
    <head>
        <meta charset='utf-8'>
        <link rel='stylesheet' href='/static/css/carta_html.css'>
    </head>
    <body>
        <h1 class='title'>OSTERIA TRE GOBBI — CARTA DEI VINI</h1>
        {body}
    </body>
    </html>
    """

    return HTMLResponse(html)


# ------------------------------------------------------------
# PDF CLIENTE
# ------------------------------------------------------------
@router.get("/carta/pdf")
def genera_carta_vini_pdf():
    data_oggi = datetime.now().strftime("%d/%m/%Y")
    rows = list(load_vini_ordinati())

    frontespizio = f"""
    <div class="front-page">
        <img src="file://{LOGO_PATH}" class="front-logo">
        <div class="front-title">CARTA VINI</div>
        <div class="front-subtitle">Aggiornata al {data_oggi}</div>
    </div>
    """

    toc_html = build_carta_toc_html(rows)
    body_html = build_carta_body_html(rows)
    carta = f"<div class='carta-body'>{body_html}</div>"

    html = f"""
    <html>
    <head>
        <meta charset='utf-8'>
        <link rel='stylesheet' href='/static/css/carta_pdf.css'>
    </head>
    <body>
        {frontespizio}
        {toc_html}
        {carta}
    </body>
    </html>
    """

    out = STATIC_DIR / "carta_vini.pdf"

    HTML(string=html, base_url=str(BASE_DIR)).write_pdf(
        str(out),
        stylesheets=[CSS(filename=str(CSS_PDF))],
    )

    return FileResponse(out, filename="carta_vini.pdf")


# ------------------------------------------------------------
# PDF STAFF
# ------------------------------------------------------------
@router.get("/carta/pdf-staff")
def genera_carta_vini_pdf_staff():
    """
    Versione STAFF.
    Per ora identica al PDF cliente, ma con label 'VERSIONE STAFF' nel frontespizio.
    """
    data_oggi = datetime.now().strftime("%d/%m/%Y")
    rows = list(load_vini_ordinati())

    frontespizio = f"""
    <div class="front-page">
        <img src="file://{LOGO_PATH}" class="front-logo">
        <div class="front-title">CARTA VINI — STAFF</div>
        <div class="front-subtitle">VERSIONE INTERNA</div>
        <div class="front-date">Aggiornata al {data_oggi}</div>
    </div>
    """

    toc_html = build_carta_toc_html(rows)
    body_html = build_carta_body_html(rows)
    carta = f"<div class='carta-body'>{body_html}</div>"

    html = f"""
    <html>
    <head>
        <meta charset='utf-8'>
        <link rel='stylesheet' href='/static/css/carta_pdf.css'>
    </head>
    <body>
        {frontespizio}
        {toc_html}
        {carta}
    </body>
    </html>
    """

    out = STATIC_DIR / "carta_vini_staff.pdf"

    HTML(string=html, base_url=str(BASE_DIR)).write_pdf(
        str(out),
        stylesheets=[CSS(filename=str(CSS_PDF))],
    )

    return FileResponse(out, filename="carta_vini_staff.pdf")


# ------------------------------------------------------------
# DOCX SEMPLICE
# ------------------------------------------------------------
@router.get("/carta/docx")
def genera_carta_vini_docx():
    rows = list(load_vini_ordinati())
    doc = Document()

    if LOGO_PATH.exists():
        doc.add_picture(str(LOGO_PATH), width=Inches(1.8))

    doc.add_heading("CARTA DEI VINI — OSTERIA TRE GOBBI", level=1)

    def k_tip(r): return r["TIPOLOGIA"] or "Senza tipologia"
    def k_reg(r): return resolve_regione(r)
    def k_prod(r): return r["PRODUTTORE"] or "Produttore sconosciuto"

    for tip, g1 in groupby(rows, k_tip):
        g1 = list(g1)
        doc.add_heading(tip, level=2)

        for reg, g2 in groupby(g1, k_reg):
            g2 = list(g2)
            doc.add_heading(reg, level=3)

            for prod, g3 in groupby(g2, k_prod):
                g3 = list(g3)
                p = doc.add_paragraph()
                r = p.add_run(prod)
                r.bold = True

                for riga in g3:
                    desc = riga["DESCRIZIONE"] or ""
                    annata = riga["ANNATA"] or ""
                    prezzo = riga["PREZZO"]
                    if prezzo:
                        try:
                            prezzo = f"€ {float(prezzo):.2f}"
                        except Exception:
                            pass

                    line = "    " + desc
                    if annata:
                        line += f" — {annata}"
                    if prezzo:
                        line += f" — {prezzo}"

                    doc.add_paragraph(line)

    out = STATIC_DIR / "carta_vini.docx"
    doc.save(str(out))

    return FileResponse(out, filename="carta_vini.docx")


# ============================================================
# MOVIMENTI CANTINA
# ============================================================
class MovimentoCreate(BaseModel):
    tipo: str = Field(..., description="CARICO / SCARICO / VENDITA / RETTIFICA")
    qta: int = Field(..., gt=0, description="Quantità positiva")
    note: str | None = Field(None, description="Nota operativa (evento, servizio, ecc.)")
    origine: str | None = Field("GESTIONALE", description="Origine movimento")
    data_mov: str | None = Field(
        None,
        description="Data movimento ISO8601; se None usa datetime.now()",
    )


@router.get("/{vino_id}/movimenti", response_class=JSONResponse)
def lista_movimenti_vino(
    vino_id: int,
    current_user: Dict[str, Any] = Depends(get_current_user),
):
    """
    Restituisce la lista dei movimenti registrati per un dato vino.
    Richiede utente loggato (JWT valido).
    """
    vino = get_vino_by_id(vino_id)
    if not vino:
        raise HTTPException(status_code=404, detail="Vino non trovato")

    rows = list_movimenti_vino(vino_id)
    movimenti = [dict(r) for r in rows]

    return JSONResponse(content={"vino_id": vino_id, "movimenti": movimenti})


@router.post("/{vino_id}/movimenti", response_class=JSONResponse)
def crea_movimento_vino(
    vino_id: int,
    payload: MovimentoCreate,
    current_user: Dict[str, Any] = Depends(get_current_user),
):
    """
    Registra un movimento di cantina per un vino:
    - CARICO    → QTA aumenta
    - SCARICO   → QTA diminuisce
    - VENDITA   → QTA diminuisce
    - RETTIFICA → qta = nuovo valore assoluto

    Logga SEMPRE l'utente che esegue il movimento, prefisso nella nota:
        [username] testo nota
    """
    vino = get_vino_by_id(vino_id)
    if not vino:
        raise HTTPException(status_code=404, detail="Vino non trovato")

    tipo = payload.tipo.strip().upper()

    # Ricavo un identificativo utente leggibile
    if isinstance(current_user, dict):
        utente = (
            current_user.get("username")
            or current_user.get("email")
            or (str(current_user.get("id")) if current_user.get("id") is not None else None)
            or "unknown"
        )
    else:
        utente = "unknown"

    # Nota completa con utente
    if payload.note:
        nota_completa = f"[{utente}] {payload.note}"
    else:
        nota_completa = f"[{utente}]"

    try:
        registra_movimento(
            vino_id=vino_id,
            tipo=tipo,
            qta=payload.qta,
            note=nota_completa,
            origine=payload.origine or "GESTIONALE",
            data_mov=payload.data_mov,
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

    return JSONResponse(
        content={
            "status": "ok",
            "vino_id": vino_id,
            "tipo": tipo,
            "qta": payload.qta,
            "eseguito_da": utente,
        }
    )